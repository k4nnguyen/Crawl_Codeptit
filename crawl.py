import sys
import io

# Ép output ra UTF-8 (để Java đọc được tiếng Việt)
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Nhận username và password từ tham số
if len(sys.argv) < 3:
    print("Vui lòng nhập username và password")
    sys.exit(1)

USERNAME = sys.argv[1]
PASSWORD = sys.argv[2]

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options

from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ---------------------
# Chủ đề cần lọc
TOPICS = [
    ("168", "LẬP TRÌNH JAVA CƠ BẢN"),
    ("169", "MẢNG"),
    ("170", "XÂU KÝ TỰ"),
    ("171", "KHAI BÁO LỚP VÀ ĐỐI TƯỢNG"),
    ("172", "MẢNG ĐỐI TƯỢNG"),
    ("173", "QUAN HỆ GIỮA CÁC LỚP"),
    ("174", "ỨNG DỤNG VÀO JAVA COLLECTION"),
    ("175", "VÀO RA FILE")
]

# ---------------------
# Khởi tạo Selenium
options = Options()
options.add_argument("--start-maximized")
driver = webdriver.Chrome(options=options)
wait = WebDriverWait(driver, 10)

# ---------------------
# 1) Đăng nhập
driver.get("https://code.ptit.edu.vn/student/question/")
wait.until(EC.presence_of_element_located((By.ID, "login__user"))).send_keys(USERNAME)
driver.find_element(By.ID, "login__pw").send_keys(PASSWORD)
driver.find_element(By.XPATH, "//button[span[text()='Đăng nhập']]").click()
wait.until(EC.presence_of_element_located((By.XPATH, "//button[i[contains(@class,'fa-filter')]]")))

# ---------------------
# 2) Tạo Word + Insert Table
doc = Document()
doc.add_heading("Danh sách bài Code PTIT", 0)

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

# set width cột
col_widths = [Inches(0.8), Inches(1.8), Inches(1.8), Inches(3.0), Inches(1.0)]
for col, w in zip(table.columns, col_widths):
    for cell in col.cells:
        cell.width = w

# tiêu đề
headers = ["STT", "Chủ đề", "Mã bài Code PTIT", "Tên bài", "Trạng thái"]
hdr_cells = table.rows[0].cells
for i, text in enumerate(headers):
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ---------------------
# 3) Lặp qua từng chủ đề
for topic_id, topic_name in TOPICS:
    print(f"Đang lọc chủ đề: {topic_name}")

    # Mở menu filter
    filter_button = wait.until(EC.element_to_be_clickable((By.XPATH, "//button[i[contains(@class,'fa-filter')]]")))
    filter_button.click()

    wait.until(EC.presence_of_all_elements_located((By.CSS_SELECTOR, "input[type='checkbox'][name='sub_group[]']")))

    # Bỏ tick tất cả
    for cb in driver.find_elements(By.CSS_SELECTOR, "input[type='checkbox'][name='sub_group[]']"):
        if cb.is_selected():
            driver.execute_script("arguments[0].click();", cb)

    # Tick chủ đề hiện tại
    cb_topic = driver.find_element(By.ID, topic_id)
    driver.execute_script("arguments[0].click();", cb_topic)

    # Bấm Lọc
    loc_button = driver.find_element(By.CLASS_NAME, "filter__dropdown__btn")
    driver.execute_script("arguments[0].click();", loc_button)

    # Chờ bảng tải
    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "tbody tr")))

    rows = driver.find_elements(By.CSS_SELECTOR, "tbody tr")
    for row in rows:
        tds = row.find_elements(By.TAG_NAME, "td")
        if len(tds) < 6:
            continue

        stt     = tds[1].text.strip()
        ma_bai  = tds[2].text.strip()
        ten_bai = tds[3].text.strip()
        chu_de  = tds[5].text.strip()

        # Xác định trạng thái
        if "bg--10th" in row.get_attribute("class"):
            trang_thai = "AC"
        else:
            trang_thai = "WA"

        # Thêm dòng vào bảng
        cells = table.add_row().cells
        values = [stt, chu_de, ma_bai, ten_bai]
        for i, val in enumerate(values):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 3 else WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(val)

        # Cột trạng thái
        p_status = cells[4].paragraphs[0]
        p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_status.add_run(trang_thai)

        # Tô nền ô trạng thái
        shade_color = "9FE2BF" if trang_thai == "AC" else "FF7F7F"
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shade_color))
        cells[4]._tc.get_or_add_tcPr().append(shading_elm)

# ---------------------
# 4) Lưu file
doc.save("Danh_sach_Code_PTIT.docx")
print("✅ Xuất file Word thành công: Danh_sach_Code_PTIT.docx")

driver.quit()
